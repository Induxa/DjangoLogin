import os
from django.conf import settings
from django.core.mail import send_mail
import datetime
import zipfile
from tempfile import TemporaryDirectory
from django.conf import settings
from docx import Document
from django.http import Http404
from django.shortcuts import render, redirect
from django.urls import reverse
from Group.models import GroupUniversity, Discip
from .models import FacultyUniversity, Specialization
from .forms import GroupDateStudent, AddFile
from openpyxl import Workbook, load_workbook
from students.models import Students, Work
from django.http import HttpResponse


def today_work(faculty_id):
    faculty = FacultyUniversity.objects.get(FacultyLoginPasswords=faculty_id)
    specialization = Specialization.objects.filter(SpecializationFaculty=faculty)
    group = GroupUniversity.objects.filter(SpecializationGroup__in=specialization).order_by('NameGroup')
    date_now = datetime.date.today()

    try:
        for groups in group:
            date_times_course = Discip.objects.filter(groups=groups, types_work='Курсовая')
            date_times_practice = Discip.objects.filter(groups=groups, types_work='Практика')
            students = Students.objects.filter(GroupStudents=groups)

            for student in students:
                send_course_mail = False
                send_practice_mail = False
                course_message = ""
                practice_message = ""

                for date_time_cr in date_times_course:
                    result_date_course = (date_time_cr.Times_files - date_now).days
                    if not result_date_course < 0 and result_date_course != 0:
                        if result_date_course <= 6 and not Work.objects.filter(DiscipId=date_time_cr, StudentId=student,
                                                                               Type='Курсовая').exists():
                            send_course_mail = True
                            course_message += f" по предмету: {date_time_cr.Name}, осталось {result_date_course} дней до сдачи курсовой."

                for date_time_pr in date_times_practice:
                    result_date_practice = (date_time_pr.Times_files - date_now).days
                    if not result_date_practice < 0 and result_date_practice != 0:
                        if result_date_practice <= 6 and not Work.objects.filter(DiscipId=date_time_pr,
                                                                                 StudentId=student,
                                                                                 Type='Практика').exists():
                            send_practice_mail = True
                            practice_message += f" и что осталось {result_date_practice} дней до сдачи практики."

                if send_course_mail:
                    send_mail(
                        f'Напоминание о сдаче курсовой работы',
                        f'Добрый день {student.Last_name} {student.First_name} {student.Middle_name}.'
                        f' студент группы: {groups.NameGroup} напоминаем вам о предстоящей сдачи курсовой {course_message}',
                        settings.EMAIL_HOST_USER,
                        [student.EmailStudents]
                    )

                if send_practice_mail:
                    send_mail(
                        f'Напоминание о сдаче практики',
                        f'Добрый день {student.Last_name} {student.First_name} {student.Middle_name}.'
                        f' студент группы: {groups.NameGroup} напоминаем вам о предстоящей сдачи практики {practice_message}',
                        settings.EMAIL_HOST_USER,
                        [student.EmailStudents]
                    )

    except Exception as ex:
        return str(ex)


def faculty_index(request, faculty_id):
    today_work(faculty_id)
    faculty = FacultyUniversity.objects.get(FacultyLoginPasswords=faculty_id)
    specialization = Specialization.objects.filter(SpecializationFaculty=faculty)
    group = GroupUniversity.objects.filter(SpecializationGroup__in=specialization).order_by('NameGroup')
    students = Students.objects.filter(GroupStudents__in=group).order_by('GroupStudents', 'Last_name')
    if request.method == "POST":
        st = request.POST.get('student_work')
        # group_students = request.POST.get('groupStudents')
        group_date = request.POST.get('select')
        student_find = request.POST.get('student_fullname')
        data_all = request.POST.get('get_data')
        if group_date:
            students = Students.objects.filter(GroupStudents__NameGroup=group_date).order_by(
                'GroupStudents',
                'Last_name')
        elif student_find:
            students = Students.objects.filter(Last_name=student_find)
        elif data_all:
            students = Students.objects.filter(GroupStudents__in=group).order_by(
                'GroupStudents', 'Last_name')
        elif st:
            return redirect(reverse('faculty:Work', kwargs={'faculty_id': faculty_id,
                                                            'students_id': st}))

    group_date = GroupDateStudent()
    data = {
        'faculty_id': faculty_id,
        'id': faculty,
        'files': AddFile,
        'group': group,
        'student': students,
        'groups': group_date,
    }
    return render(request, 'faculty/index.html', context=data)


def graph(request, faculty_id):
    faculty = FacultyUniversity.objects.get(FacultyLoginPasswords=faculty_id)
    specialization = Specialization.objects.filter(SpecializationFaculty=faculty)
    group = GroupUniversity.objects.filter(SpecializationGroup__in=specialization).order_by('NameGroup')
    date_times = Discip.objects.filter(groups__in=group).order_by('groups')

    if request.method == 'POST':
        for key, value in request.POST.items():
            if key.startswith('group_'):
                group_id = key.split('_')[0]
                group_name = value
                courses = Discip.objects.filter(groups__NameGroup=group_name, types_work='Курсовая')
                practices = Discip.objects.filter(groups__NameGroup=group_name, types_work='Практика')
            if key.startswith('dateCourse_'):
                date_id = key.split('_')[0]
                date_course = value
            if key.startswith('datePractice_'):
                date_id = key.split('_')[0]
                date_practice = value
                for cours in courses:
                    if cours.Times_files != date_course and date_course != '':
                        course = Discip.objects.filter(groups__NameGroup=group_name, types_work='Курсовая').update(
                            Times_files=date_course)
                    elif cours.Times_files == date_course and date_course == '':
                        continue
                for pract in practices:
                    if pract.Times_files != date_practice and date_practice != '':
                        practice = Discip.objects.filter(groups__NameGroup=group_name, types_work='Практика').update(
                            Times_files=date_practice)
                    elif pract.Times_files == date_practice and date_practice == '':
                        continue

        return redirect(reverse('faculty:graph', kwargs={'faculty_id': faculty_id}))

    grouped_times = {}
    for date_time in date_times:
        group_name = date_time.groups.NameGroup
        if group_name not in grouped_times:
            grouped_times[group_name] = {'Курсовая': None, 'Практика': None}
        if date_time.types_work == 'Курсовая':
            grouped_times[group_name]['Курсовая'] = date_time.Times_files
        elif date_time.types_work == 'Практика':
            grouped_times[group_name]['Практика'] = date_time.Times_files

    data = {
        'faculty_id': faculty_id,
        'groups': group,
        'grouped_times': grouped_times
    }

    return render(request, 'faculty/graph.html', context=data)


def logoutUser(request):
    # logout(request)
    pass


def download_file(request):
    if request.method == "POST":
        download = request.POST.get('download')

        if not download:
            raise Http404("No file specified for download.")

        file_extension = os.path.splitext(download)[1].lower()
        content_type = 'application/octet-stream'  # Default content type

        if file_extension == '.zip':
            content_type = 'application/zip'
        elif file_extension == '.rar':
            content_type = 'application/x-rar-compressed'

        file_path = os.path.join(settings.MEDIA_ROOT, download)

        if os.path.exists(file_path):
            with open(file_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type=content_type)
                response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
                return response
        else:
            raise Http404("File does not exist.")

    raise Http404("Invalid request method.")


def Working(request, faculty_id, students_id):
    students = Students.objects.filter(pk=students_id)
    faculty = FacultyUniversity.objects.get(FacultyLoginPasswords=faculty_id)
    specialization = Specialization.objects.filter(SpecializationFaculty=faculty)
    if Work.objects.filter(StudentId=students_id).exists():
        themes = Work.objects.filter(StudentId=students_id)
    elif not Work.objects.filter(StudentId=students_id).exists():
        themes = Work.objects.filter(StudentId=students_id)

    data = {
        'faculty_id': faculty_id,
        'students': students,
        'themes': themes,

    }
    return render(request, 'faculty/Student_work.html', context=data)


def Parsing(request, faculty_id):
    if request.method == "POST":
        files = AddFile(request.POST, request.FILES)
        if files.is_valid():
            uploaded_file = request.FILES['Files']
            file_name = uploaded_file.name
            if file_name.endswith('.docx'):
                doc = Document(uploaded_file)
                for groups in doc.paragraphs:
                    if groups.text.strip():
                        group = groups.text.strip()
                        name, code = group.split('-')
                for table in doc.tables:
                    for row in table.rows:
                        fio = row.cells[0].text.strip()
                        fio_parts = fio.split()
                        last_name = fio_parts[0]
                        theme = row.cells[1].text.strip()
                        student = Students.objects.get(Last_name=last_name)
                        name_themes = Work.objects.filter(StudentId=student.id, Course=code[0],
                                                          Type='Курсовая')
                        if name_themes.exists():
                            for update in name_themes:
                                update.Name_themes = theme
                                update.save()

    return redirect(reverse('faculty:faculty_index', kwargs={'faculty_id': faculty_id}))


def archive(request, faculty_id):
    faculty = FacultyUniversity.objects.get(FacultyLoginPasswords=faculty_id)
    specialization = Specialization.objects.filter(SpecializationFaculty=faculty)
    group = GroupUniversity.objects.filter(SpecializationGroup__in=specialization)

    if request.method == "POST":
        group_name = request.POST.get('select')
        student_find = request.POST.get('student_fullname')
        semester = request.POST.get('course')
        if group_name:
            group_id = GroupUniversity.objects.get(NameGroup=group_name)
            student_fullname = Students.objects.filter(GroupStudents=group_id).order_by('Last_name')
            files_to_zip = []
            with TemporaryDirectory() as temp_dir:
                for st in student_fullname:
                    student_id = Students.objects.get(Last_name=st.Last_name)
                    id_student = student_id.id
                    course_data = Work.objects.filter(StudentId=id_student, Type='Курсовая')
                    pracite_data = Work.objects.filter(StudentId=id_student, Type='Практика')
                    for course in course_data:

                        course_dir = os.path.join(temp_dir, f"Курс {course.Course}")
                        os.makedirs(course_dir, exist_ok=True)

                        folder_course = os.path.join(course_dir, 'Курсовые')
                        os.makedirs(folder_course, exist_ok=True)

                        course_folder = os.path.join(folder_course,
                                                     f'{course.StudentId.Last_name} '
                                                     f'{course.StudentId.First_name[0]}. '
                                                     f'{course.StudentId.Middle_name[0]}.')

                        os.makedirs(course_folder, exist_ok=True)

                        if course.Course:
                            file_path = os.path.join(course_folder, os.path.basename(course.Files.name))
                            with open(file_path, 'wb') as f:
                                f.write(course.Files.read())
                                files_to_zip.append(file_path)
                    for practice in pracite_data:

                        practice_dir = os.path.join(temp_dir, f"Курс {practice.Course}")
                        os.makedirs(practice_dir, exist_ok=True)

                        folder_practice = os.path.join(practice_dir, 'Практики')
                        os.makedirs(folder_practice, exist_ok=True)

                        parc = os.path.join(folder_practice, f'{practice.StudentId.Last_name} '
                                                             f'{practice.StudentId.First_name[0]}. '
                                                             f'{practice.StudentId.Middle_name[0]}.')
                        os.makedirs(parc, exist_ok=True)

                        if practice.Course:
                            file_path = os.path.join(parc, os.path.basename(practice.Files.name))
                            with open(file_path, 'wb') as f:
                                f.write(practice.Files.read())
                                files_to_zip.append(file_path)

                zip_file_path = os.path.join(temp_dir, 'IT-11.zip')
                with zipfile.ZipFile(zip_file_path, 'w') as zipf:
                    for file_path in files_to_zip:
                        zipf.write(file_path, os.path.relpath(file_path, temp_dir))

                with open(zip_file_path, 'rb') as zip_file:
                    response = HttpResponse(zip_file.read(), content_type='application/zip')
                    response['Content-Disposition'] = 'inline; filename=' + group_name
                    return response

        if student_find:
            student_fullname = Students.objects.get(Last_name=student_find)
            student_id = student_fullname.id
            course_data = Work.objects.filter(StudentId=student_id, Type='Курсовая')
            pracite_data = Work.objects.filter(StudentId=student_id, Type='Практика')

            with TemporaryDirectory() as temp_dir:

                for course in course_data:
                    course_dir = os.path.join(temp_dir, f"Курс {course.Course}")
                    os.makedirs(course_dir, exist_ok=True)

                    folder_course = os.path.join(course_dir, 'Курсовая')
                    os.makedirs(folder_course, exist_ok=True)

                    if course.Course:
                        file_path = os.path.join(folder_course, os.path.basename(course.Files.name))
                        with open(file_path, 'wb') as f:
                            f.write(course.Files.read())

                for pracite in pracite_data:
                    course_dir = os.path.join(temp_dir, f"Курс {pracite.Course}")
                    os.makedirs(course_dir, exist_ok=True)

                    folder_practice = os.path.join(course_dir, 'Практика')
                    os.makedirs(folder_practice, exist_ok=True)

                    if pracite.Course:
                        file_path = os.path.join(folder_practice, os.path.basename(pracite.Files.name))
                        with open(file_path, 'wb') as f:
                            f.write(pracite.Files.read())

                zip_file_path = os.path.join(temp_dir, 'IT-11.zip')
                with zipfile.ZipFile(zip_file_path, 'w') as zipf:
                    for root, dirs, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            zipf.write(file_path, os.path.relpath(file_path, temp_dir))

                with open(zip_file_path, 'rb') as zip_file:
                    response = HttpResponse(zip_file.read(), content_type='application/zip')
                    response['Content-Disposition'] = f'attachment; filename="{student_find}.zip"'
                    return response

    data = {
        'faculty_id': faculty_id,
        'id': faculty,
        'groups': group,
    }
    return render(request, 'faculty/archive.html', context=data)


def faculty_students(request, faculty_id):
    faculty = FacultyUniversity.objects.get(FacultyLoginPasswords=faculty_id)
    specialization = Specialization.objects.filter(SpecializationFaculty=faculty)
    group = GroupUniversity.objects.filter(SpecializationGroup__in=specialization)
    students = Students.objects.all().order_by('Last_name')
    # if request.method == "POST":
    #     status = request.POST.get('status')
    #     work = request.POST.get('student_work')
    #     if status == '1':
    #         students = Students.objects.filter(Status_Students=1).order_by('Last_name')
    #         # return redirect(reverse('faculty:faculty_students', kwargs={'faculty_id': faculty_id,
    #         #                                                             'students': students}))
    #     elif status == '2':
    #         students = Students.objects.filter(GroupStudents__in=group, Status_Students=2).order_by('Last_name')
    #     #             return redirect(reverse('faculty:faculty_students', kwargs={'faculty_id': faculty_id,
    #     #                                                                         'students': students}))
    #     elif status == '3':
    #         students = Students.objects.filter(GroupStudents__in=group, Status_Students=3).order_by('Last_name')
    #     #             return redirect(reverse('faculty:faculty_students', kwargs={'faculty_id': faculty_id,
    #     #                                                                         'students': students}))
    #     elif status == '4':
    #         students = Students.objects.filter(GroupStudents__in=group, Status_Students=4).order_by('Last_name')
    #     #             return redirect(reverse('faculty:faculty_students', kwargs={'faculty_id': faculty_id,
    #     #                                                                         'students': students}))
    #     elif status == '5':
    #         students = Students.objects.filter(Status_Students=5).order_by('Last_name')
    #     #             return redirect(reverse('faculty:faculty_students', kwargs={'faculty_id': faculty_id,
    #     #                                                                         'students': students}))
    #     elif work:
    #         return redirect(reverse('faculty:Work', kwargs={'faculty_id': faculty_id,
    #                                                         'students_id': work}))
    data = {
        'faculty_id': faculty_id,
        'students': students,
    }
    return render(request, 'faculty/faculty_students.html', context=data)


def Report(request, faculty_id):
    faculty = FacultyUniversity.objects.get(FacultyLoginPasswords=faculty_id)
    specialization = Specialization.objects.filter(SpecializationFaculty=faculty)
    group = GroupUniversity.objects.filter(SpecializationGroup__in=specialization).order_by('NameGroup')

    if request.method == 'POST':
        name_group = request.POST.get('select')

        students = Students.objects.filter(GroupStudents__NameGroup=name_group).order_by('GroupStudents', 'Last_name')
        for gp in group:
            group_name = gp.NameGroup
        wb = Workbook()
        ws = wb.active

        ws.merge_cells('A1:A2')
        ws['A1'] = 'ФИО'
        ws.merge_cells('B1:C1')
        ws['B1'] = 'Курсовая 1 курс'
        ws['B2'] = 'листы'
        ws['C2'] = 'электроно'
        ws.merge_cells('D1:E1')
        ws['D1'] = 'Курсовая 2 курс'
        ws['D2'] = 'листы'
        ws['E2'] = 'электроно'
        ws.merge_cells('F1:G1')
        ws['F1'] = 'Курсовая 3 курс'
        ws['F2'] = 'листы'
        ws['G2'] = 'электроно'
        ws.merge_cells('H1:I1')
        ws['H1'] = 'Курсовая 4 курс'
        ws['H2'] = 'листы'
        ws['I2'] = 'электроно'
        ws.merge_cells('J1:K1')
        ws['J1'] = 'Практика 1 курс'
        ws['J2'] = 'листы'
        ws['K2'] = 'электроно'
        ws.merge_cells('L1:M1')
        ws['L1'] = 'Практика 2 курс'
        ws['L2'] = 'листы'
        ws['M2'] = 'электроно'
        ws.merge_cells('N1:O1')
        ws['N1'] = 'Практика 3 курс'
        ws['N2'] = 'листы'
        ws['O2'] = 'электроно'
        ws.merge_cells('P1:Q1')
        ws['P1'] = 'Практика 4 курс'
        ws['P2'] = 'листы'
        ws['Q2'] = 'электроно'

        for row_index, st in enumerate(students, start=3):
            ws.cell(row=row_index, column=1).value = f'{st.Last_name} {st.First_name[0]}. {st.Middle_name[0]}.'
            course_data = Work.objects.filter(StudentId=st.id, Type='Курсовая')
            practice_data = Work.objects.filter(StudentId=st.id, Type='Практика')
            if course_data.exists():
                for course in course_data:
                    if course.Course == '1':
                        ws.cell(row=row_index, column=3).value = '+'
                    elif course.Course == '2':
                        ws.cell(row=row_index, column=5).value = '+'
                    elif course.Course == '3':
                        ws.cell(row=row_index, column=7).value = '+'
                    elif course.Course == '4':
                        ws.cell(row=row_index, column=9).value = '+'
            if practice_data.exists():
                for practice in practice_data:
                    if practice.Course == '1':
                        ws.cell(row=row_index, column=11).value = '+'
                    if practice.Course == '2':
                        ws.cell(row=row_index, column=13).value = '+'
                    if practice.Course == '3':
                        ws.cell(row=row_index, column=15).value = '+'
                    if practice.Course == '4':
                        ws.cell(row=row_index, column=17).value = '+'
            row_index += 1

        ws.title = f'{group_name}'
        file_name = f'Отчёт по {group_name}.xlsx'
        file_path = os.path.join('C:\\Users\\barka\\Desktop\\DjangoLogin\\', file_name)
        wb.save(file_path)

        with open(file_path, 'rb') as excel_file:
            response = HttpResponse(excel_file.read(),
                                    content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response['Content-Disposition'] = f'attachment; filename="{file_name}"'

        return response


def transfer(request, faculty_id):
    faculty = FacultyUniversity.objects.get(FacultyLoginPasswords=faculty_id)
    specialization = Specialization.objects.filter(SpecializationFaculty=faculty)
    group = GroupUniversity.objects.filter(SpecializationGroup__in=specialization)
    namegroup = []
    for gp in group:
        name = gp.NameGroup.split('-')[0]
        namegroup.append(name)
    group_name = None
    students = None
    if request.method == 'POST':
        gt = request.POST.get('group_view')
        group_find = request.POST.get('group_find')
        if gt:
            group_name = GroupUniversity.objects.filter(SpecializationGroup__in=specialization, NameGroup__contains=gt)
        elif group_find:
            students = Students.objects.filter(GroupStudents__NameGroup=group_find).exclude(Status_Students=2).exclude(
                Status_Students=5).order_by(
                'GroupStudents',
                'Last_name')
        for key, value in request.POST.items():
            if key.startswith('st_'):
                student = key.split('_')[0]
                st = value
                last_name = st.split()[0]
                student = Students.objects.get(Last_name=last_name)
            if key.startswith('status_'):
                student_id = key.split('_')[0]
                status = value
                if student:
                    student_group = student.GroupStudents
                    print(f'new status: {status}')
                    student.Status_Students = status
                    student.save()
                    group_number = int(student_group.NameGroup.split('-')[1])
                    if not '2' in str(student.Status_Students):
                        if group.exists():
                            if '4' in str(group_number):
                                student.Status_Students = '5'
                                student.save()
                            elif '25' in str(group_number):
                                student.Status_Students = '5'
                                student.save()
                            else:
                                group_text, group_number2 = student_group.NameGroup.split('-')
                                group_number3 = int(group_number2) + 10
                                next_group = f'{group_text}-{group_number3}'
                                create_new_group = GroupUniversity.objects.filter(
                                    SpecializationGroup__in=specialization).filter(
                                    NameGroup=next_group)
                                if create_new_group.exists():
                                    group_go = GroupUniversity.objects.get(NameGroup=next_group)
                                    student.GroupStudents = group_go
                                    student.save()
                                elif not create_new_group.exists():
                                    specialization_get = Specialization.objects.get(SpecializationFaculty=faculty)
                                    create_group = GroupUniversity.objects.create(NameGroup=next_group,
                                                                                  SpecializationGroup=specialization_get)
                                    create_group.NameGroup = next_group
                                    create_group.save()
                                    return redirect(reverse('faculty:transfer', kwargs={'faculty_id': faculty_id}))
    data = {
        'id': faculty,
        'faculty_id': faculty_id,
        'groups': group_name,
        'name_group': sorted(set(namegroup)),
        'students': students
    }
    return render(request, 'faculty/transfer_students.html', context=data)
